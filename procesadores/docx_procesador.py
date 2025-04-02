import os
import docx
from .base_procesador import ProcesadorDocumento

class ProcesadorDOCX(ProcesadorDocumento):
    """Procesador para archivos Word (.docx)"""
    
    def extraer_texto(self, ruta_archivo):
        """
        Extrae el texto de un archivo .docx con marcadores de página
        
        Args:
            ruta_archivo (str): Ruta al archivo a procesar
            
        Returns:
            str: Texto extraído del documento con marcadores de página
        """
        try:
            doc = docx.Document(ruta_archivo)
            texto_completo = []
            num_pagina = 1
            parrafos_por_pagina = 15  # Estimación aproximada de párrafos por página
            contador_parrafos = 0
            
            # Marcar el inicio del documento con la primera página
            texto_completo.append(f"[PAGINA:{num_pagina}]")
            
            # Extraer texto de los párrafos con estimación de cambios de página
            for para in doc.paragraphs:
                # Agregar el texto del párrafo
                if para.text.strip():  # Solo considerar párrafos con contenido
                    texto_completo.append(para.text)
                    contador_parrafos += 1
                    
                    # Estimar cambios de página basados en la cantidad de párrafos
                    # También considerar párrafos largos como posibles cambios de página
                    if contador_parrafos >= parrafos_por_pagina or len(para.text) > 500:
                        num_pagina += 1
                        texto_completo.append(f"\n[PAGINA:{num_pagina}]")
                        contador_parrafos = 0
            
            # Extraer texto de las tablas (las tablas generalmente ocupan más espacio)
            for tabla in doc.tables:
                num_pagina += 1  # Asumir que cada tabla comienza en una nueva página
                texto_completo.append(f"\n[PAGINA:{num_pagina}]")
                
                for fila in tabla.rows:
                    fila_texto = []
                    for celda in fila.cells:
                        if celda.text.strip():
                            fila_texto.append(celda.text.strip())
                    if fila_texto:
                        texto_completo.append(" | ".join(fila_texto))
                
                contador_parrafos = 0
            
            return '\n'.join(texto_completo)
        except Exception as e:
            print(f"Error al procesar el DOCX {ruta_archivo}: {e}")
            return ""
    
    def obtener_metadatos(self, ruta_archivo):
        """
        Obtiene los metadatos del archivo .docx
        
        Args:
            ruta_archivo (str): Ruta al archivo a procesar
            
        Returns:
            dict: Metadatos del documento
        """
        metadatos = {
            'nombre': os.path.basename(ruta_archivo),
            'extension': '.docx',
            'tamaño': os.stat(ruta_archivo).st_size,
            'fecha_modificacion': os.stat(ruta_archivo).st_mtime,
            'ruta': ruta_archivo
        }
        
        try:
            doc = docx.Document(ruta_archivo)
            
            # Extraer propiedades básicas
            core_props = doc.core_properties
            if core_props:
                if core_props.title:
                    metadatos['titulo'] = core_props.title
                if core_props.author:
                    metadatos['autor'] = core_props.author
                if core_props.created:
                    metadatos['fecha_creacion'] = core_props.created
                if core_props.modified:
                    metadatos['fecha_modificacion_doc'] = core_props.modified
                if core_props.comments:
                    metadatos['comentarios'] = core_props.comments
                
        except Exception as e:
            print(f"Error al obtener metadatos del DOCX {ruta_archivo}: {e}")
            
        return metadatos 